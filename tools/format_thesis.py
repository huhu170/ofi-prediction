"""
format_thesis.py — 上海财经大学 MBA 论文 Word 排版后处理脚本
读取 pandoc 生成的 paper_raw.docx，按格式要求应用样式，输出 paper_final.docx
"""
import sys
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------
FONT_CN = "宋体"
FONT_EN = "Times New Roman"
FONT_HEADING_CN = "黑体"
FONT_QUOTE = "楷体"
FONT_CODE = "Consolas"

PT_TITLE = Pt(16)        # 论文标题
PT_CHAPTER = Pt(14)      # 章标题 (Heading 1)
PT_SECTION = Pt(12)      # 节标题 (Heading 2)
PT_SUBSEC = Pt(12)       # 目标题 (Heading 3)
PT_BODY = Pt(10.5)       # 正文 (五号)
PT_TABLE = Pt(9)         # 表格内容
PT_TABLE_NOTE = Pt(9)    # 表注
PT_CAPTION = Pt(10.5)    # 图表标题
PT_REF = Pt(10.5)        # 参考文献

LINE_BODY = Pt(20)       # 正文固定行距 20 磅
LINE_REF = Pt(16)        # 参考文献行距

FIRST_LINE_INDENT = Pt(21)  # 约2字符首行缩进

TABLE_HEADER_BG = "D9E2F3"  # 淡蓝灰
BORDER_COLOR = "000000"

REF_RE = re.compile(r"^\s*\[?\d+\]?")
REF_STRICT_RE = re.compile(r"^\s*\\\[?\d+\\\]?|^\s*\[\d+\]")
EQ_NUM_RE = re.compile(r"^（[\d\.\-a-z]+）$")
NOTE_PREFIXES = ("注：", "注:", "数据来源", "说明：", "说明:", "表注：", "表注:")


# ---------------------------------------------------------------------------
# 辅助函数
# ---------------------------------------------------------------------------

def set_run_font(run, size, bold=False, italic=False, color=None,
                 cn_font=FONT_CN, en_font=FONT_EN):
    """设置 run 的中英文字体、字号、加粗、颜色"""
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)


def set_spacing(para, before=None, after=None, line=None, rule=None):
    """设置段落间距与行距"""
    pf = para.paragraph_format
    if before is not None:
        pf.space_before = before
    if after is not None:
        pf.space_after = after
    if line is not None:
        pf.line_spacing = line
    if rule is not None:
        pf.line_spacing_rule = rule


def para_has_math(para):
    """段落是否包含 OMML 公式"""
    return bool(
        para._element.findall(".//" + qn("m:oMath"))
        or para._element.findall(".//" + qn("m:oMathPara"))
    )


def para_has_image(para):
    """段落是否包含图片"""
    return bool(para._element.findall(".//" + qn("w:drawing")))


def set_cell_border(cell, val="single", sz="4", color=BORDER_COLOR):
    """设置单元格四边边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(borders)
    for edge in ("top", "bottom", "left", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" '
                f'w:space="0" w:color="{color}"/>'
            )
            borders.append(el)
        else:
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), sz)
            el.set(qn("w:color"), color)


def set_cell_shading(cell, color_hex):
    """设置单元格底纹"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(doc):
    """在页脚添加居中页码"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run()
        run._r.append(parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
        ))
        run2 = p.add_run()
        run2._r.append(parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        ))
        run3 = p.add_run()
        run3._r.append(parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        ))
        for r in (run, run2, run3):
            set_run_font(r, PT_BODY)


# ---------------------------------------------------------------------------
# 分类判断
# ---------------------------------------------------------------------------

def is_title_para(para, idx):
    """论文标题（第0段，非 Heading 样式）"""
    return idx == 0 and not (para.style and para.style.name.startswith("Heading"))


def is_abstract_heading(para):
    """摘要 / Abstract 标题"""
    txt = para.text.strip()
    sn = para.style.name if para.style else ""
    return sn.startswith("Heading") and txt in ("摘要", "Abstract")


def is_keyword_para(para):
    """关键词段"""
    txt = para.text.strip()
    return txt.startswith("关键词") or txt.startswith("Keywords")


def is_table_note(para):
    """表注段落"""
    txt = para.text.strip()
    return any(txt.startswith(p) for p in NOTE_PREFIXES)


def is_caption_like(para):
    """图表标题：Image Caption 样式 或 全bold且以'图/表'+编号开头的短段落"""
    sn = para.style.name if para.style else ""
    if sn == "Image Caption":
        return True
    txt = para.text.strip()
    if not re.match(r"^[图表]\s*\d", txt):
        return False
    runs = para.runs
    if not runs:
        return False
    all_bold = all(r.bold for r in runs if r.text.strip())
    if all_bold and len(txt) < 100:
        return True
    return False


# ---------------------------------------------------------------------------
# 格式化各类段落
# ---------------------------------------------------------------------------

def fmt_title(para):
    """论文大标题：黑体 16pt 居中 加粗"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(para, Pt(24), Pt(18), Pt(28), WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        set_run_font(run, PT_TITLE, bold=True, cn_font=FONT_HEADING_CN)


def fmt_heading1(para):
    """章标题：黑体 14pt 居中 加粗"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(para, Pt(18), Pt(12), None, WD_LINE_SPACING.SINGLE)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        set_run_font(run, PT_CHAPTER, bold=True, cn_font=FONT_HEADING_CN)


def fmt_heading2(para, center=False):
    """节标题：黑体 12pt"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, Pt(12), Pt(6), None, WD_LINE_SPACING.SINGLE)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        set_run_font(run, PT_SECTION, bold=True, cn_font=FONT_HEADING_CN)


def fmt_heading3(para):
    """目标题：黑体 12pt 左对齐"""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, Pt(6), Pt(3), None, WD_LINE_SPACING.SINGLE)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        set_run_font(run, PT_SUBSEC, bold=True, cn_font=FONT_HEADING_CN)


def fmt_body(para):
    """正文段落：宋体/TNR 10.5pt 首行缩进 固定行距20磅"""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, Pt(0), Pt(0), LINE_BODY, WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    for run in para.runs:
        was_bold = run.bold
        set_run_font(run, PT_BODY, bold=bool(was_bold))


def fmt_equation(para):
    """公式段落：居中 前后各6磅间距 无缩进"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(para, Pt(6), Pt(6), LINE_BODY, WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None
    for run in para.runs:
        set_run_font(run, PT_BODY)


def fmt_image(para):
    """图片段落：居中 无缩进"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(para, Pt(6), Pt(3))
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None


def fmt_caption(para):
    """图表标题：居中 10.5pt 前后各6磅"""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(para, Pt(3), Pt(6), LINE_BODY, WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None
    for run in para.runs:
        set_run_font(run, PT_CAPTION, bold=bool(run.bold))


def fmt_table_note(para):
    """表注：9pt 宋体 左对齐 无缩进 行距固定16磅"""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, Pt(2), Pt(2), Pt(16), WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None
    for run in para.runs:
        set_run_font(run, PT_TABLE_NOTE)


def fmt_keyword(para):
    """关键词段：无缩进 10.5pt"""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, Pt(6), Pt(12), LINE_BODY, WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        was_bold = run.bold
        set_run_font(run, PT_BODY, bold=bool(was_bold))


def fmt_eq_number(para):
    """公式编号：右对齐 10.5pt 无缩进"""
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(para, Pt(0), Pt(6), LINE_BODY, WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None
    for run in para.runs:
        set_run_font(run, PT_BODY)


def fmt_ref_entry(para):
    """参考文献条目：10.5pt 行距16磅 悬挂缩进"""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, Pt(0), Pt(0), Pt(16), WD_LINE_SPACING.EXACTLY)
    pf = para.paragraph_format
    pf.left_indent = Pt(21)
    pf.first_line_indent = -Pt(21)
    for run in para.runs:
        set_run_font(run, PT_REF)


def fmt_block_text(para):
    """引用块：楷体 10.5pt 左缩进1cm"""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, Pt(3), Pt(3), LINE_BODY, WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        set_run_font(run, PT_BODY, cn_font=FONT_QUOTE)


def fmt_code(para):
    """代码块：等宽字体 9pt 左对齐 无缩进"""
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(para, Pt(0), Pt(0), Pt(14), WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = Cm(0.5)
    for run in para.runs:
        set_run_font(run, Pt(9), cn_font=FONT_CODE, en_font=FONT_CODE)


def fmt_list_item(para):
    """列表项：宋体 10.5pt 左缩进 无首行缩进"""
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(para, Pt(0), Pt(0), LINE_BODY, WD_LINE_SPACING.EXACTLY)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        was_bold = run.bold
        set_run_font(run, PT_BODY, bold=bool(was_bold))


def is_primarily_math(para):
    """段落是否以公式为主体（非行内公式的正文）"""
    if not para_has_math(para):
        return False
    plain = para.text.strip()
    return len(plain) < 20 or not plain


# ---------------------------------------------------------------------------
# 主格式化流程
# ---------------------------------------------------------------------------

def format_document(doc):
    in_ref_section = False

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        sn = para.style.name if para.style else ""

        # --- 论文大标题 ---
        if is_title_para(para, idx):
            fmt_title(para)
            continue

        # --- 标题样式 ---
        if sn.startswith("Heading"):
            if text.startswith("参考文献"):
                in_ref_section = True
            if sn.startswith("Heading 1"):
                fmt_heading1(para)
            elif sn.startswith("Heading 2"):
                center = is_abstract_heading(para) or text.startswith("参考文献")
                fmt_heading2(para, center=center)
            elif sn.startswith("Heading 3"):
                fmt_heading3(para)
            continue

        # --- 参考文献区域 ---
        if in_ref_section:
            if REF_RE.match(text) or text[:1].isdigit() or text.startswith("["):
                fmt_ref_entry(para)
            elif text:
                fmt_body(para)
            continue

        # --- 代码块（必须在正文之前） ---
        if sn == "Source Code":
            fmt_code(para)
            continue

        # --- 关键词 ---
        if is_keyword_para(para):
            fmt_keyword(para)
            continue

        # --- 表注/注释（必须在公式检查之前，避免含行内公式的注释被误判为公式） ---
        if is_table_note(para):
            fmt_table_note(para)
            continue

        # --- 引用块 Block Text（同样必须在公式检查之前）---
        if sn == "Block Text":
            if any(text.startswith(p) for p in NOTE_PREFIXES):
                fmt_table_note(para)
            elif ("†" in text or "‡" in text
                    or text.startswith("脚本") or "脚本：" in text):
                fmt_table_note(para)
            else:
                fmt_block_text(para)
            continue

        # --- 图片段落 ---
        if para_has_image(para) or sn == "Captioned Figure":
            fmt_image(para)
            continue

        # --- 图表标题 ---
        if is_caption_like(para):
            fmt_caption(para)
            continue

        # --- 公式段落（仅对以公式为主体的段落居中） ---
        if is_primarily_math(para):
            fmt_equation(para)
            continue

        # --- 公式编号（如 "（3.2-1）"）---
        if EQ_NUM_RE.match(text):
            fmt_eq_number(para)
            continue

        # --- 列表项 ---
        if sn == "Compact":
            fmt_list_item(para)
            continue

        # --- 空段落 / 分隔段 ---
        if not text:
            set_spacing(para, Pt(0), Pt(0), Pt(6), WD_LINE_SPACING.EXACTLY)
            para.paragraph_format.first_line_indent = None
            continue

        # --- 默认：正文（含 First Paragraph、Body Text 等） ---
        fmt_body(para)


def format_tables(doc):
    """格式化所有表格"""
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_border(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if row_idx == 0:
                    set_cell_shading(cell, TABLE_HEADER_BG)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_spacing(para, Pt(1), Pt(1))
                    para.paragraph_format.first_line_indent = None
                    para.paragraph_format.left_indent = None
                    for run in para.runs:
                        bold = (row_idx == 0) or bool(run.bold)
                        set_run_font(run, PT_TABLE, bold=bold)


def format_page(doc):
    """页面设置"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)
    add_page_number(doc)


# ---------------------------------------------------------------------------
# 入口
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 3:
        print("Usage: python format_thesis.py <input.docx> <output.docx>")
        sys.exit(1)

    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])

    print(f"Loading: {src}")
    doc = Document(str(src))

    print("  [1/4] Formatting paragraphs...")
    format_document(doc)

    print("  [2/4] Formatting tables...")
    format_tables(doc)

    print("  [3/4] Setting page layout...")
    format_page(doc)

    print("  [4/4] Saving...")
    doc.save(str(dst))
    print(f"Done: {dst}")


if __name__ == "__main__":
    main()
