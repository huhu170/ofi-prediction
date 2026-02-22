"""Quick verification of paper_final.docx formatting."""
import re
from docx import Document

doc = Document(r"d:\paper project\paper_final.docx")

s = doc.sections[0]
print("=== Page Layout ===")
print(f"  Size: {s.page_width.cm:.1f} x {s.page_height.cm:.1f} cm")
print(f"  Margins: T={s.top_margin.cm:.2f} B={s.bottom_margin.cm:.2f} "
      f"L={s.left_margin.cm:.2f} R={s.right_margin.cm:.2f}")

print("\n=== Sample Headings ===")
for p in doc.paragraphs[:50]:
    if p.style.name.startswith("Heading"):
        runs_info = ""
        if p.runs:
            r = p.runs[0]
            runs_info = f"size={r.font.size}, bold={r.font.bold}"
        print(f'  [{p.style.name}] "{p.text[:40]}" => {runs_info}')

print("\n=== Sample Body Text ===")
body_count = 0
for p in doc.paragraphs:
    if p.style.name == "Normal" and len(p.text.strip()) > 20 and body_count < 3:
        pf = p.paragraph_format
        print(f'  "{p.text[:50]}..."')
        print(f"    line_spacing={pf.line_spacing}, first_indent={pf.first_line_indent}")
        if p.runs:
            r = p.runs[0]
            print(f"    font_size={r.font.size}, font_name={r.font.name}")
        body_count += 1

print(f"\n=== Tables: {len(doc.tables)} total ===")
if doc.tables:
    t = doc.tables[0]
    print(f"  Table 0: {len(t.rows)} rows x {len(t.columns)} cols")
    if t.rows and t.rows[0].cells:
        c = t.rows[0].cells[0]
        if c.paragraphs and c.paragraphs[0].runs:
            r = c.paragraphs[0].runs[0]
            print(f"  Header cell font: size={r.font.size}, bold={r.font.bold}")

print("\n=== Reference Entries (sample) ===")
ref_count = 0
for p in doc.paragraphs:
    if re.match(r"^\[\d+\]", p.text.strip()) and ref_count < 3:
        pf = p.paragraph_format
        print(f'  "{p.text[:60]}..."')
        print(f"    line_spacing={pf.line_spacing}, first_indent={pf.first_line_indent}")
        ref_count += 1

print("\n=== Footer ===")
for section in doc.sections:
    f = section.footer
    for p in f.paragraphs:
        print(f'  Footer text: "{p.text}", alignment={p.alignment}')
