"""
md2docx.py — 论文 Markdown → Word 完整转换脚本
1. 预处理 Markdown（处理 \\tag{}、修复图片路径）
2. 调用 pandoc 生成初始 docx（嵌入图片 + 公式转 OMML）
3. python-docx 后处理排版（标题/正文/表格/公式间距/参考文献/页面）
"""
import os
import re
import sys
import subprocess
import tempfile
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ===================================================================
# 配置常量
# ===================================================================
FONT_CN = "宋体"
FONT_EN = "Times New Roman"
FONT_HEADING_CN = "黑体"
FONT_QUOTE = "楷体"

PT_TITLE = Pt(22)         # 论文总标题
PT_CHAPTER = Pt(16)       # 章标题 (Heading 1)
PT_SECTION = Pt(14)       # 节标题 (Heading 2)
PT_SUBSEC = Pt(12)        # 目标题 (Heading 3)
PT_BODY = Pt(12)          # 正文 小四号 ≈ 12pt
PT_TABLE = Pt(10.5)       # 表格 五号
PT_REF = Pt(10.5)         # 参考文献 五号
PT_CAPTION = Pt(10.5)     # 图表标题 五号
PT_FOOTER = Pt(10.5)      # 页码

LINE_BODY_PT = 20         # 正文固定行距 20磅
LINE_REF_PT = 20          # 参考文献行距

TABLE_HEADER_BG = "D9E2F3"  # 表头浅蓝
BORDER_COLOR = "000000"

REF_RE = re.compile(r"^\s*\[\d+\]")
CN_CHAR_RE = re.compile(r"[\u4e00-\u9fff]")

# ===================================================================
# 第一步：预处理 Markdown
# ===================================================================

def preprocess_markdown(md_text: str, project_dir: str) -> str:
    """
    预处理 Markdown 文本：
    1. 将 \\tag{...} 提取为公式编号，放在公式块后作为文本
    2. 修复图片路径中的 %20（URL编码空格 → 实际空格）
    """
    lines = md_text.split("\n")
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # 处理行内图片链接中的 %20
        if re.match(r"^!\[", line):
            line = line.replace("%20", " ")
            # 验证图片文件是否存在
            match = re.search(r'\]\((.+?)\)', line)
            if match:
                img_path = match.group(1)
                full_path = os.path.join(project_dir, img_path)
                if not os.path.exists(full_path):
                    print(f"  警告: 图片不存在 → {full_path}")
            result.append(line)
            i += 1
            continue

        # 处理 $$ ... $$ 公式块（可能跨行或单行）
        if line.strip().startswith("$$"):
            eq_lines = [line]
            # 检查是否单行 $$ ... $$
            if line.strip().endswith("$$") and line.strip() != "$$":
                # 单行公式
                eq_text = line.strip()
                tag_num = _extract_tag(eq_text)
                eq_text = _remove_tag(eq_text)
                result.append(eq_text)
                if tag_num:
                    result.append("")
                    result.append(f"<p style='text-align:right'>({tag_num})</p>")
                i += 1
                continue

            # 多行公式：寻找结束 $$
            i += 1
            while i < len(lines):
                eq_lines.append(lines[i])
                if lines[i].strip().endswith("$$"):
                    break
                i += 1
            i += 1

            # 合并公式行
            eq_block = "\n".join(eq_lines)
            tag_num = _extract_tag(eq_block)
            eq_block = _remove_tag(eq_block)
            result.append(eq_block)
            if tag_num:
                result.append("")
                result.append(f"<p style='text-align:right'>({tag_num})</p>")
            continue

        result.append(line)
        i += 1

    return "\n".join(result)


def _extract_tag(text: str) -> str:
    """从公式中提取 \\tag{...} 的编号"""
    match = re.search(r"\\tag\{([^}]+)\}", text)
    return match.group(1) if match else ""


def _remove_tag(text: str) -> str:
    """移除 \\tag{...}"""
    return re.sub(r"\s*\\tag\{[^}]+\}", "", text)


# ===================================================================
# 第二步：调用 pandoc 转换
# ===================================================================

def run_pandoc(md_path: str, docx_path: str, project_dir: str):
    """调用 pandoc 将 Markdown 转为 docx"""
    # 在 Windows 上刷新 PATH
    pandoc_paths = [
        r"C:\Users\12995\AppData\Local\Microsoft\WinGet\Packages\JohnMacFarlane.Pandoc_Microsoft.Winget.Source_8wekyb3d8bbwe\pandoc-3.9\pandoc.exe",
        r"C:\Program Files\Pandoc\pandoc.exe",
        r"C:\Users\12995\AppData\Local\Pandoc\pandoc.exe",
    ]
    pandoc_cmd = "pandoc"
    for p in pandoc_paths:
        if os.path.exists(p):
            pandoc_cmd = p
            break

    cmd = [
        pandoc_cmd,
        str(md_path),
        "-o", str(docx_path),
        "--from", "markdown+tex_math_dollars+raw_html",
        "--resource-path", str(project_dir),
        "--standalone",
    ]
    print(f"  运行 pandoc ...")
    result = subprocess.run(cmd, capture_output=True, text=True,
                            cwd=project_dir, encoding="utf-8",
                            errors="replace")
    if result.returncode != 0:
        print(f"  pandoc 错误: {result.stderr}")
        sys.exit(1)
    if result.stderr:
        # 打印警告但不退出
        warnings = result.stderr.strip().split("\n")
        for w in warnings[:10]:
            print(f"  pandoc 警告: {w}")
        if len(warnings) > 10:
            print(f"  ... 共 {len(warnings)} 条警告")
    print(f"  pandoc 完成 → {docx_path}")


# ===================================================================
# 第三步：python-docx 后处理
# ===================================================================

def set_run_font(run, size, bold=False, italic=False, color=None,
                 cn_font=None, en_font=None):
    """设置 run 的中英文字体"""
    _cn = cn_font or FONT_CN
    _en = en_font or FONT_EN
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), _cn)
    rfonts.set(qn("w:ascii"), _en)
    rfonts.set(qn("w:hAnsi"), _en)


def set_paragraph_fmt(para, align=None, space_before=None, space_after=None,
                      line_spacing_pt=None, first_indent=None, left_indent=None,
                      hanging=None):
    """统一设置段落格式"""
    pf = para.paragraph_format
    if align is not None:
        para.alignment = align
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing_pt is not None:
        pf.line_spacing = Pt(line_spacing_pt)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if hanging is not None:
        pf.first_line_indent = -hanging
        pf.left_indent = hanging


def set_cell_borders(cell, val="single", sz="4", color=BORDER_COLOR):
    """设置单元格四边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(borders)
    for edge in ("top", "bottom", "left", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is not None:
            borders.remove(el)
        new_el = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        )
        borders.append(new_el)


def set_cell_shading(cell, color_hex):
    """设置单元格底纹"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def is_equation_para(para):
    """判断段落是否为公式段落（包含 OMML 或 oMath 元素）"""
    xml = para._p.xml
    if "m:oMath" in xml or "m:oMathPara" in xml:
        return True
    return False


def is_equation_number_para(para):
    """判断是否为公式编号段落，如 (3.1-2)"""
    text = para.text.strip()
    return bool(re.match(r"^\(\d+\.\d+[-–]\d+[a-z]?\)$", text))


def is_caption_para(para):
    """判断是否为图表标题"""
    text = para.text.strip()
    if not text:
        return False
    return (text.startswith("图 ") or text.startswith("图\u00a0")
            or text.startswith("表 ") or text.startswith("表\u00a0")
            or text.startswith("图1") or text.startswith("图2")
            or text.startswith("图3") or text.startswith("图4")
            or text.startswith("表1") or text.startswith("表2")
            or text.startswith("表3") or text.startswith("表4"))


def is_image_para(para):
    """判断段落是否包含图片"""
    for run in para.runs:
        if run._r.findall(qn("w:drawing")):
            return True
        if run._r.findall(qn("w:pict")):
            return True
    xml_str = para._p.xml
    return "w:drawing" in xml_str or "w:pict" in xml_str


def format_document(doc):
    """主格式化入口"""
    _format_page_layout(doc)
    _format_all_paragraphs(doc)
    _format_tables(doc)
    _add_page_numbers(doc)


def _format_page_layout(doc):
    """A4页面、标准边距"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)


def _format_all_paragraphs(doc):
    """遍历所有段落并应用格式"""
    in_ref_section = False
    is_title = True  # 第一个非空段落视为标题

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        # 跳过空段落
        if not text and not is_image_para(para) and not is_equation_para(para):
            continue

        # 论文总标题（第一个非空段落）
        if is_title and text and not style_name.startswith("Heading"):
            is_title = False
            _fmt_title(para)
            continue
        if is_title and style_name.startswith("Heading"):
            is_title = False

        # 检测参考文献节
        if text.startswith("参考文献") and style_name.startswith("Heading"):
            in_ref_section = True
            _fmt_heading(para, 1)
            continue

        # 标题
        if style_name.startswith("Heading 1") or style_name == "Heading 1":
            _fmt_heading(para, 1)
            continue
        elif style_name.startswith("Heading 2") or style_name == "Heading 2":
            _fmt_heading(para, 2)
            continue
        elif style_name.startswith("Heading 3") or style_name == "Heading 3":
            _fmt_heading(para, 3)
            continue

        # 参考文献条目
        if in_ref_section and REF_RE.match(text):
            _fmt_ref_entry(para)
            continue

        # 公式段落
        if is_equation_para(para):
            _fmt_equation(para)
            continue

        # 公式编号
        if is_equation_number_para(para):
            _fmt_equation_number(para)
            continue

        # 图片段落
        if is_image_para(para):
            _fmt_image(para)
            continue

        # 图表标题
        if is_caption_para(para):
            _fmt_caption(para)
            continue

        # 普通正文
        _fmt_body(para)


def _fmt_title(para):
    """论文总标题：黑体二号居中加粗"""
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=Pt(24), space_after=Pt(18),
                      line_spacing_pt=28)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        set_run_font(run, PT_TITLE, bold=True, cn_font=FONT_HEADING_CN)


def _fmt_heading(para, level):
    """格式化标题"""
    sizes = {1: PT_CHAPTER, 2: PT_SECTION, 3: PT_SUBSEC}
    size = sizes.get(level, PT_SUBSEC)

    if level == 1:
        align = WD_ALIGN_PARAGRAPH.CENTER
        sb, sa = Pt(24), Pt(12)
        ls = 28
    elif level == 2:
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        sb, sa = Pt(18), Pt(6)
        ls = 24
    else:
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        sb, sa = Pt(12), Pt(6)
        ls = 22

    set_paragraph_fmt(para, align=align, space_before=sb, space_after=sa,
                      line_spacing_pt=ls)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        set_run_font(run, size, bold=True, cn_font=FONT_HEADING_CN)


def _fmt_body(para):
    """正文：宋体小四号(12pt)，首行缩进2字符，行距20磅"""
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      space_before=Pt(0), space_after=Pt(0),
                      line_spacing_pt=LINE_BODY_PT,
                      first_indent=Cm(0.85))  # 约2字符 @ 12pt
    for run in para.runs:
        is_bold = run.bold
        is_italic = run.italic
        set_run_font(run, PT_BODY, bold=bool(is_bold), italic=bool(is_italic))


def _fmt_ref_entry(para):
    """参考文献条目：五号，悬挂缩进，行距20磅"""
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      space_before=Pt(0), space_after=Pt(2),
                      line_spacing_pt=LINE_REF_PT,
                      hanging=Cm(0.85))
    for run in para.runs:
        set_run_font(run, PT_REF)


def _fmt_equation(para):
    """公式段落：居中，前后留间距"""
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=Pt(6), space_after=Pt(6),
                      line_spacing_pt=24)
    para.paragraph_format.first_line_indent = None


def _fmt_equation_number(para):
    """公式编号：右对齐"""
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.RIGHT,
                      space_before=Pt(0), space_after=Pt(6),
                      line_spacing_pt=LINE_BODY_PT)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        set_run_font(run, PT_BODY)


def _fmt_image(para):
    """图片段落：居中"""
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=Pt(6), space_after=Pt(3))
    para.paragraph_format.first_line_indent = None


def _fmt_caption(para):
    """图表标题：居中，五号"""
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=Pt(3), space_after=Pt(6),
                      line_spacing_pt=LINE_BODY_PT)
    para.paragraph_format.first_line_indent = None
    for run in para.runs:
        set_run_font(run, PT_CAPTION, bold=bool(run.bold))


def _format_tables(doc):
    """格式化所有表格"""
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_borders(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if row_idx == 0:
                    set_cell_shading(cell, TABLE_HEADER_BG)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_fmt(para, space_before=Pt(1), space_after=Pt(1),
                                      line_spacing_pt=16)
                    para.paragraph_format.first_line_indent = None
                    for run in para.runs:
                        bold = (row_idx == 0) or bool(run.bold)
                        set_run_font(run, PT_TABLE, bold=bold)


def _add_page_numbers(doc):
    """页脚居中页码"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        run1 = p.add_run()
        run1._r.append(parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
        ))
        run2 = p.add_run()
        run2._r.append(parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        ))
        run3 = p.add_run()
        run3._r.append(parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        ))
        for r in (run1, run2, run3):
            set_run_font(r, PT_FOOTER)


# ===================================================================
# 图片尺寸调整
# ===================================================================

def resize_images(doc):
    """将过大的图片缩放到页面宽度以内"""
    max_width = Cm(14.66)  # A4 - 左右边距 (21 - 3.17*2)
    for para in doc.paragraphs:
        for run in para.runs:
            drawings = run._r.findall(qn("w:drawing"))
            for drawing in drawings:
                # 查找 extent 元素
                for ext in drawing.iter():
                    if ext.tag.endswith("}extent") or ext.tag == "extent":
                        cx = int(ext.get("cx", 0))
                        cy = int(ext.get("cy", 0))
                        if cx > max_width:
                            ratio = max_width / cx
                            ext.set("cx", str(int(max_width)))
                            ext.set("cy", str(int(cy * ratio)))
                    # Also check a:ext inside a:xfrm
                    if "xfrm" in ext.tag:
                        for child in ext:
                            if child.tag.endswith("}ext"):
                                cx = int(child.get("cx", 0))
                                cy = int(child.get("cy", 0))
                                if cx > max_width:
                                    ratio = max_width / cx
                                    child.set("cx", str(int(max_width)))
                                    child.set("cy", str(int(cy * ratio)))


# ===================================================================
# 主入口
# ===================================================================

def main():
    project_dir = Path(__file__).resolve().parent.parent
    md_path = project_dir / "paper.md"
    raw_docx = project_dir / "paper_raw.docx"
    final_docx = project_dir / "paper_final_v2.docx"

    print("=" * 60)
    print("论文 Markdown → Word 转换")
    print("=" * 60)

    # 1. 读取并预处理 Markdown
    print("\n[1/4] 预处理 Markdown ...")
    md_text = md_path.read_text(encoding="utf-8")
    preprocessed = preprocess_markdown(md_text, str(project_dir))

    # 写入临时文件
    prep_path = project_dir / "paper_prep.md"
    prep_path.write_text(preprocessed, encoding="utf-8")
    print(f"  预处理完成 → {prep_path}")

    # 2. 调用 pandoc
    print("\n[2/4] 调用 pandoc 转换 ...")
    run_pandoc(str(prep_path), str(raw_docx), str(project_dir))

    # 3. 后处理格式
    print("\n[3/4] 后处理排版格式 ...")
    doc = Document(str(raw_docx))
    format_document(doc)
    resize_images(doc)

    # 4. 保存
    print("\n[4/4] 保存最终文档 ...")
    doc.save(str(final_docx))
    print(f"\n完成！最终文档: {final_docx}")

    # 清理临时文件
    if prep_path.exists():
        prep_path.unlink()
        print(f"  已清理临时文件: {prep_path}")


if __name__ == "__main__":
    main()
